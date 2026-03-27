"""
Document Processor
==================
Unified document processor that orchestrates OCR and intelligent extraction
for maximum accuracy in CV/Resume processing.
"""

import logging
import base64
import json
import re
from typing import Optional, Dict, Any, Tuple, Union
from pathlib import Path
from dataclasses import dataclass, asdict
from datetime import datetime

logger = logging.getLogger(__name__)

from app.cv_parser.pdf_parser import pdf_parser, PYMUPDF_AVAILABLE
from app.cv_parser.ocr_engine import get_ocr_engine, OCREngine
from app.cv_parser.intelligent_extractor import (
    get_intelligent_extractor, 
    IntelligentCVExtractor,
    ExtractedCVData
)

# Check for python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available for Word document parsing")


@dataclass
class ProcessingResult:
    """Result of document processing."""
    success: bool
    extracted_data: Optional[ExtractedCVData] = None
    raw_text: Optional[str] = None
    text_source: str = "unknown"  # pdf_text, ocr_tesseract, ocr_openai, docx
    text_confidence: float = 0.0
    extraction_confidence: float = 0.0
    error_message: Optional[str] = None
    warnings: list = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        result = asdict(self)
        if self.extracted_data:
            result['extracted_data'] = self.extracted_data.to_dict()
        return result


class DocumentProcessor:
    """
    Unified document processor for CV/Resume files.
    
    Features:
    - Automatic format detection (PDF, Word, Image)
    - Intelligent text extraction (embedded text vs OCR)
    - LLM-powered structured data extraction
    - Confidence scoring at each stage
    - Fallback mechanisms for reliability
    """
    
    # Supported file extensions
    SUPPORTED_PDF = ['.pdf']
    SUPPORTED_WORD = ['.doc', '.docx']
    SUPPORTED_IMAGE = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']
    SUPPORTED_TEXT = ['.txt', '.rtf']
    
    def __init__(
        self,
        ocr_engine: Optional[OCREngine] = None,
        intelligent_extractor: Optional[IntelligentCVExtractor] = None
    ):
        """
        Initialize the document processor.
        
        Args:
            ocr_engine: OCR engine instance
            intelligent_extractor: Intelligent extractor instance
        """
        self.ocr_engine = ocr_engine
        self.intelligent_extractor = intelligent_extractor
    
    def _ensure_engines(self):
        """Ensure OCR and extraction engines are initialized."""
        if self.ocr_engine is None:
            self.ocr_engine = get_ocr_engine()
        if self.intelligent_extractor is None:
            self.intelligent_extractor = get_intelligent_extractor()

    def _clean_and_parse_json(self, raw_text: str) -> Dict[str, Any]:
        """Aggressively strip markdown wrappers and parse JSON safely."""
        try:
            clean_text = (raw_text or "").strip()

            if clean_text.startswith("```json"):
                clean_text = clean_text[7:]
            elif clean_text.startswith("```"):
                clean_text = clean_text[3:]
            if clean_text.endswith("```"):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            # Capture the largest JSON object body when extra prose is present.
            match = re.search(r"\{.*\}", clean_text, re.DOTALL)
            if match:
                clean_text = match.group(0)

            return json.loads(clean_text)
        except json.JSONDecodeError as exc:
            logger.error(f"CV JSON parsing failed: {exc}. Raw content: {raw_text!r}")
            return {
                "name": "Unknown",
                "phone": None,
                "experience_years": None,
                "skills": [],
                "extraction_failed": True,
            }
    
    def process_document(
        self,
        file_content: bytes,
        filename: str,
        use_intelligent_extraction: bool = True,
        use_openai_ocr: bool = True,
        expected_language: str = 'en',
        image_url: Optional[str] = None,
    ) -> ProcessingResult:
        """
        Process a CV/Resume document with maximum accuracy.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            use_intelligent_extraction: Whether to use LLM for extraction
            use_openai_ocr: Whether to use OpenAI Vision for OCR
            expected_language: Expected language of the document
            
        Returns:
            ProcessingResult with extracted data
        """
        self._ensure_engines()
        
        ext = Path(filename).suffix.lower()
        
        try:
            if ext in self.SUPPORTED_IMAGE and use_intelligent_extraction and image_url:
                vision_extracted = self._extract_structured_from_image_url(image_url)
                if vision_extracted:
                    warnings = vision_extracted.warnings.copy()
                    if vision_extracted.missing_critical_fields:
                        warnings.append(
                            f"Missing required fields: {', '.join(vision_extracted.missing_critical_fields)}"
                        )
                    return ProcessingResult(
                        success=True,
                        extracted_data=vision_extracted,
                        raw_text=vision_extracted.raw_text,
                        text_source="vision_gpt4o_json",
                        text_confidence=0.92,
                        extraction_confidence=vision_extracted.overall_confidence,
                        warnings=warnings,
                    )

            # Step 1: Extract raw text based on file type
            if ext in self.SUPPORTED_PDF:
                text, text_source, text_confidence = self._process_pdf(
                    file_content, use_openai_ocr, expected_language
                )
            elif ext in self.SUPPORTED_WORD:
                text, text_source, text_confidence = self._process_word(file_content)
            elif ext in self.SUPPORTED_IMAGE:
                text, text_source, text_confidence = self._process_image(
                    file_content, use_openai_ocr, expected_language
                )
            elif ext in self.SUPPORTED_TEXT:
                text, text_source, text_confidence = self._process_text(file_content)
            else:
                return ProcessingResult(
                    success=False,
                    error_message=f"Unsupported file format: {ext}"
                )
            
            if not text or len(text.strip()) < 20:
                return ProcessingResult(
                    success=False,
                    error_message="Could not extract text from document. Please ensure the file is readable.",
                    raw_text=text,
                    text_source=text_source
                )
            
            # Step 2: Extract structured data using intelligent extraction
            if use_intelligent_extraction:
                extracted_data = self.intelligent_extractor.extract_from_text(text)
            else:
                # Use basic regex extraction
                from app.cv_parser.text_extractor import text_extractor, CVData
                basic_data = text_extractor.extract_from_text(text)
                extracted_data = self._convert_basic_to_extracted(basic_data)
            
            # Calculate overall extraction confidence
            extraction_confidence = extracted_data.overall_confidence
            
            # Collect warnings
            warnings = extracted_data.warnings.copy()
            
            if text_confidence < 0.7:
                warnings.append(f"Text extraction confidence is low ({text_confidence:.1%})")
            
            if extracted_data.missing_critical_fields:
                warnings.append(f"Missing required fields: {', '.join(extracted_data.missing_critical_fields)}")
            
            return ProcessingResult(
                success=True,
                extracted_data=extracted_data,
                raw_text=text,
                text_source=text_source,
                text_confidence=text_confidence,
                extraction_confidence=extraction_confidence,
                warnings=warnings
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return ProcessingResult(
                success=False,
                error_message=f"Processing error: {str(e)}"
            )

    def _extract_structured_from_image_url(self, image_url: str) -> Optional[ExtractedCVData]:
        """
        Route CV photos directly to GPT-4o Vision and request structured JSON.
        """
        if not self.intelligent_extractor or not getattr(self.intelligent_extractor, "openai_client", None):
            return None

        try:
            response = self.intelligent_extractor.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert CV parser for recruitment workflows. Return valid JSON only.",
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Extract all recruitment data (Name, Phone, Passport, Experience, Skills) "
                                    "from this image of a CV. The text may be blurry or poorly lit. "
                                    "Be highly tolerant of typos. Return JSON with keys: "
                                    "name, phone, passport, experience, skills, raw_text, confidence, missing_critical_fields."
                                ),
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": image_url,
                                    "detail": "high",
                                },
                            },
                        ],
                    },
                ],
                response_format={"type": "json_object"},
            )

            payload = self._clean_and_parse_json(response.choices[0].message.content or "")
            name = (payload.get("name") or "").strip() or None
            phone = (payload.get("phone") or "").strip() or None
            passport = (payload.get("passport") or "").strip() or None
            experience = payload.get("experience")
            if experience is None:
                experience = payload.get("experience_years")
            skills = payload.get("skills")
            if isinstance(skills, str):
                skills = [s.strip() for s in skills.split(",") if s.strip()]
            if not isinstance(skills, list):
                skills = []

            confidence = float(payload.get("confidence", 0.8) or 0.8)
            missing_critical_fields = payload.get("missing_critical_fields")
            if not isinstance(missing_critical_fields, list):
                missing_critical_fields = []
            if not phone:
                missing_critical_fields.append("phone")
            missing_critical_fields = list(dict.fromkeys(missing_critical_fields))

            total_exp = None
            try:
                if experience is not None:
                    total_exp = float(experience)
            except Exception:
                total_exp = None

            return ExtractedCVData(
                full_name=name,
                full_name_confidence=confidence if name else 0.0,
                phone=phone,
                phone_confidence=confidence if phone else 0.0,
                expected_salary=passport,
                expected_salary_confidence=confidence if passport else 0.0,
                total_experience_years=total_exp,
                total_experience_years_confidence=confidence if total_exp is not None else 0.0,
                technical_skills=skills,
                raw_text=(payload.get("raw_text") or "").strip() or None,
                extraction_method="vision_gpt4o_json",
                extraction_timestamp=datetime.utcnow().isoformat(),
                overall_confidence=max(0.0, min(confidence, 1.0)),
                missing_critical_fields=missing_critical_fields,
                warnings=["Vision extraction used for image CV input"],
            )
        except Exception as exc:
            logger.warning(f"Vision structured extraction failed, falling back to OCR path: {exc}")
            return None
    
    def _process_pdf(
        self,
        file_content: bytes,
        use_openai_ocr: bool,
        language: str
    ) -> Tuple[str, str, float]:
        """
        Process a PDF file.
        First tries embedded text extraction, falls back to OCR if needed.
        """
        # Try extracting embedded text first
        embedded_text = pdf_parser.extract_text_from_bytes(file_content)
        embedded_len = len((embedded_text or "").strip())
        short_pdf_threshold = 50
        
        if embedded_text and embedded_len > 100:
            # Check if the text looks valid (not just artifacts)
            if self._is_valid_text(embedded_text):
                logger.info("Using embedded PDF text")
                return embedded_text, "pdf_embedded", 0.95

        # If PDF text is too short, treat as likely scanned and OCR before accepting partial text.
        if self.ocr_engine and 0 < embedded_len < short_pdf_threshold:
            logger.info(
                f"PDF embedded text too short ({embedded_len} chars). Falling back to OCR."
            )
            text, confidence = self.ocr_engine.extract_text_from_pdf_images(
                file_content, use_openai_ocr, language
            )
            if text:
                source = "ocr_openai_short_pdf" if use_openai_ocr else "ocr_tesseract_short_pdf"
                return text, source, confidence
        
        # Check if it's a scanned PDF
        if self.ocr_engine and self.ocr_engine.is_scanned_pdf(file_content):
            logger.info("Detected scanned PDF, applying OCR")
            text, confidence = self.ocr_engine.extract_text_from_pdf_images(
                file_content, use_openai_ocr, language
            )
            
            if text:
                source = "ocr_openai" if use_openai_ocr else "ocr_tesseract"
                return text, source, confidence
        
        # If we have some embedded text, use it
        if embedded_text and embedded_len > 20:
            return embedded_text, "pdf_embedded_partial", 0.6
        
        # Last resort: try full OCR even if not detected as scanned
        if self.ocr_engine:
            logger.info("Attempting OCR as fallback")
            text, confidence = self.ocr_engine.extract_text_from_pdf_images(
                file_content, use_openai_ocr, language
            )
            
            if text:
                source = "ocr_openai_fallback" if use_openai_ocr else "ocr_tesseract_fallback"
                return text, source, confidence
        
        return "", "pdf_failed", 0.0
    
    def _process_word(self, file_content: bytes) -> Tuple[str, str, float]:
        """Process a Word document."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return "", "docx_unavailable", 0.0
        
        try:
            import io
            doc = DocxDocument(io.BytesIO(file_content))
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            
            text = '\n'.join(paragraphs)
            
            if text:
                return text, "docx_extracted", 0.95
            else:
                return "", "docx_empty", 0.0
                
        except Exception as e:
            logger.error(f"Word document processing failed: {e}")
            return "", "docx_failed", 0.0
    
    def _process_image(
        self,
        file_content: bytes,
        use_openai_ocr: bool,
        language: str
    ) -> Tuple[str, str, float]:
        """Process an image file (CV photo/scan)."""
        if not self.ocr_engine:
            logger.error("OCR engine not available for image processing")
            return "", "image_ocr_unavailable", 0.0
        
        text, confidence = self.ocr_engine.extract_text_from_image(
            file_content, use_openai_ocr, language
        )
        
        source = "ocr_openai" if use_openai_ocr else "ocr_tesseract"
        return text, source, confidence
    
    def _process_text(self, file_content: bytes) -> Tuple[str, str, float]:
        """Process a plain text file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    return text, f"text_{encoding}", 0.99
                except UnicodeDecodeError:
                    continue
            
            # Fallback with error handling
            text = file_content.decode('utf-8', errors='ignore')
            return text, "text_fallback", 0.8
            
        except Exception as e:
            logger.error(f"Text file processing failed: {e}")
            return "", "text_failed", 0.0
    
    def _is_valid_text(self, text: str) -> bool:
        """
        Check if extracted text looks valid and not just artifacts.
        """
        if not text:
            return False
        
        # Count actual words
        words = text.split()
        word_count = len(words)
        
        if word_count < 10:
            return False
        
        # Check for reasonable character ratio
        alpha_chars = sum(1 for c in text if c.isalpha())
        total_chars = len(text)
        
        if total_chars == 0:
            return False
        
        alpha_ratio = alpha_chars / total_chars
        
        # Should have at least 30% alphabetic characters
        return alpha_ratio > 0.3
    
    def _convert_basic_to_extracted(self, basic_data) -> ExtractedCVData:
        """Convert basic CVData to ExtractedCVData."""
        from datetime import datetime
        
        return ExtractedCVData(
            full_name=basic_data.name,
            full_name_confidence=0.6 if basic_data.name else 0.0,
            email=basic_data.email,
            email_confidence=0.9 if basic_data.email else 0.0,
            phone=basic_data.phone,
            phone_confidence=0.8 if basic_data.phone else 0.0,
            address=basic_data.address,
            address_confidence=0.5 if basic_data.address else 0.0,
            highest_qualification=basic_data.highest_qualification,
            highest_qualification_confidence=0.5 if basic_data.highest_qualification else 0.0,
            current_job_title=basic_data.current_position,
            current_company=basic_data.current_company,
            total_experience_years=float(basic_data.experience_years) if basic_data.experience_years else None,
            notice_period=basic_data.notice_period,
            technical_skills=basic_data.skills.split(', ') if basic_data.skills else [],
            languages_spoken=basic_data.languages or [],
            profile_summary=basic_data.summary,
            raw_text=basic_data.raw_text,
            extraction_method="regex_basic",
            extraction_timestamp=datetime.utcnow().isoformat(),
            missing_critical_fields=basic_data.missing_fields or []
        )
    
    def process_whatsapp_image(
        self,
        image_bytes: bytes,
        expected_language: str = 'en',
        image_url: Optional[str] = None,
    ) -> ProcessingResult:
        """
        Process an image received from WhatsApp (photo of CV).
        Uses high-accuracy OCR for best results.
        
        Args:
            image_bytes: Image content as bytes
            expected_language: Expected language
            
        Returns:
            ProcessingResult with extracted data
        """
        return self.process_document(
            file_content=image_bytes,
            filename="whatsapp_image.jpg",
            use_intelligent_extraction=True,
            use_openai_ocr=True,  # Always use best OCR for WhatsApp images
            expected_language=expected_language,
            image_url=image_url,
        )


# Factory function
def create_document_processor() -> DocumentProcessor:
    """Create document processor with configuration."""
    return DocumentProcessor()


# Lazy-loaded singleton
_document_processor: Optional[DocumentProcessor] = None

def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor singleton."""
    global _document_processor
    if _document_processor is None:
        _document_processor = create_document_processor()
    return _document_processor
